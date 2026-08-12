from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor

from app.analysis.models.analysis import (
    AnalysisResult,
    DocumentCoverage,
    Evidence,
    KeyFinding,
    Recommendation,
)

BANK_NAME = "Meridian Bank"
CLASSIFICATION = "Internal — Confidential"
NAVY = RGBColor(0x1E, 0x3A, 0x5F)
GREY = RGBColor(0x6B, 0x72, 0x80)

_MONTHS = [
    "January", "February", "March", "April", "May", "June",
    "July", "August", "September", "October", "November", "December",
]


def _format_month(value: str | None) -> str:
    """Convert a ``YYYY-MM`` value to ``Month YYYY`` (or return the raw value)."""
    if not value:
        return ""
    try:
        year, month = value.split("-", 1)
        return f"{_MONTHS[int(month) - 1]} {year}"
    except (ValueError, IndexError):
        return value


def _format_date(value: str | None) -> str:
    """Normalise ISO timestamps to ``Aug 6, 2026`` style text."""
    if not value:
        return ""
    try:
        dt = datetime.fromisoformat(value.replace("Z", "+00:00"))
    except ValueError:
        return value
    return dt.strftime("%b %d, %Y")


def _period(result: AnalysisResult) -> str:
    if result.trend:
        return _format_month(result.trend[-1].month)
    if result.summary.generatedAt:
        return _format_date(result.summary.generatedAt)
    return _format_date(result.createdAt)


def _evidence_text(evidence: list[Evidence]) -> str:
    parts: list[str] = []
    for item in evidence:
        source = item.sourceRef or item.documentType or item.documentId
        if item.snippet:
            parts.append(f"{source}: {item.snippet}")
        else:
            parts.append(source)
    return " · ".join(parts)


def _clean(value: str | None, limit: int = 320) -> str:
    text = (value or "").replace("\n", " ").replace("\r", " ").strip()
    if len(text) > limit:
        return text[: limit - 1].rstrip() + "…"
    return text


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for index, header in enumerate(headers):
        hdr[index].text = ""
        run = hdr[index].paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = NAVY
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = _clean(value)
            for run in cells[index].paragraphs[0].runs:
                run.font.size = Pt(9)
    doc.add_paragraph()


def _add_heading(doc: Document, text: str) -> None:
    heading = doc.add_heading(text, level=1)
    for run in heading.runs:
        run.font.color.rgb = NAVY


def _metric_by_id(result: AnalysisResult, metric_id: str) -> int | None:
    for metric in result.metrics:
        if metric.id == metric_id:
            return metric.value
    return None


def _documents_of_category(result: AnalysisResult, *categories: str) -> list[DocumentCoverage]:
    return [doc for doc in result.documents if doc.category in categories]


def build_word_report(result: AnalysisResult) -> bytes:
    doc = Document()

    # Base styling.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    # --- Cover / classification block -------------------------------------------------
    classification = doc.add_paragraph()
    classification.alignment = 1  # centered
    run = classification.add_run(CLASSIFICATION)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = GREY

    title = doc.add_paragraph()
    title.alignment = 1
    run = title.add_run(BANK_NAME)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY

    subtitle = doc.add_paragraph()
    subtitle.alignment = 1
    run = subtitle.add_run("Monthly Operational Risk Report")
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY

    meta = doc.add_paragraph()
    meta.alignment = 1
    run = meta.add_run(
        f"Reporting period: {_period(result)}   ·   "
        f"Generated: {_format_date(result.summary.generatedAt or result.createdAt)}   ·   "
        f"Prepared by Risk Copilot AI"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    doc.add_paragraph()

    # --- Document metadata -------------------------------------------------------------
    _add_heading(doc, "1. Document Information")
    _add_table(
        doc,
        ["Report ID", "Status", "Confidence", "Documents Analyzed"],
        [
            [
                result.id,
                result.status,
                f"{result.confidence}%",
                str(len(result.documents)) if result.documents else "0",
            ]
        ],
    )

    # --- Executive summary -------------------------------------------------------------
    _add_heading(doc, "2. Executive Summary")
    paragraphs = result.summary.paragraphs or []
    if not paragraphs:
        paragraphs = [
            "No executive summary was generated for this analysis run."
        ]
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)

    if result.metrics:
        doc.add_paragraph("Headline metrics:").runs[0].bold = True
        _add_table(
            doc,
            ["Metric", "Value"],
            [[metric.label, str(metric.value)] for metric in result.metrics],
        )

    # --- Board-level summary -----------------------------------------------------------
    _add_heading(doc, "3. Board-Level Summary")
    score = result.overallScore
    doc.add_paragraph(
        f"The overall operational risk score is {score.score}/100 ({score.level}). "
        f"{score.description or 'See the risk score section below for detail.'}"
    )
    key_messages: list[str] = []
    if result.keyFindings:
        top = result.keyFindings[0]
        key_messages.append(
            f"Highest-ranked risk: {top.severity} — {top.title}"
        )
        critical = [f for f in result.keyFindings if f.severity == "Critical"]
        if critical:
            key_messages.append(
                f"{len(critical)} critical finding(s) feature in the top-ranked risks."
            )
    overdue_exceptions = _metric_by_id(result, "overdue-exceptions")
    if overdue_exceptions:
        key_messages.append(
            f"{overdue_exceptions} exception(s) are overdue for remediation."
        )
    overdue_findings = _metric_by_id(result, "overdue-findings")
    if overdue_findings:
        key_messages.append(
            f"{overdue_findings} audit finding(s) are overdue."
        )
    if not key_messages:
        key_messages.append("No standout risk messages to highlight in this cycle.")
    for message in key_messages:
        doc.add_paragraph(message, style="List Bullet")

    # --- Risk overview -----------------------------------------------------------------
    _add_heading(doc, "4. Risk Overview")
    distribution = [
        ("Critical", _metric_by_id(result, "critical-risks")),
        ("High", _metric_by_id(result, "high-risks")),
        ("Medium", _metric_by_id(result, "medium-risks")),
        ("Low", _metric_by_id(result, "low-risks")),
    ]
    rows = [
        [severity, str(value) if value is not None else "n/a"]
        for severity, value in distribution
    ]
    total = _metric_by_id(result, "total-risks")
    if total is not None:
        rows.append(["Total", str(total)])
    _add_table(doc, ["Severity", "Risk Count"], rows)

    if result.heatmap:
        doc.add_paragraph("Risk heatmap by division and category:").runs[0].bold = True
        heat_rows: list[list[str]] = []
        for row in result.heatmap:
            heat_rows.append(
                [
                    row.division,
                    ", ".join(
                        f"{cell.category} ({cell.level})" for cell in row.cells
                    ),
                ]
            )
        _add_table(doc, ["Division", "Category exposure"], heat_rows)

    if result.trend:
        doc.add_paragraph("Monthly risk trend:").runs[0].bold = True
        _add_table(
            doc,
            ["Month", "High", "Medium", "Low"],
            [
                [
                    _format_month(point.month),
                    str(point.high),
                    str(point.medium),
                    str(point.low),
                ]
                for point in result.trend
            ],
        )

    # --- Risk score --------------------------------------------------------------------
    _add_heading(doc, "5. Risk Score")
    doc.add_paragraph(f"Overall risk score: {score.score}/100 ({score.level})")
    if score.change:
        doc.add_paragraph(f"Change: {score.change}")
    doc.add_paragraph(score.description or "No description available.")
    doc.add_paragraph(f"Analysis confidence: {result.confidence}%")

    # --- Key findings ------------------------------------------------------------------
    _add_heading(doc, "6. Key Findings")
    findings = result.keyFindings or []
    if not findings:
        doc.add_paragraph("No key findings were generated for this analysis run.")
    else:
        _add_table(
            doc,
            ["ID", "Finding", "Category", "Severity", "Likelihood", "Exposure", "Confidence", "Evidence"],
            [
                [
                    finding.id,
                    finding.title,
                    finding.category,
                    finding.severity,
                    finding.likelihood,
                    finding.exposure,
                    f"{finding.confidence}%",
                    _evidence_text(finding.evidence),
                ]
                for finding in findings
            ],
        )

    # --- Audit findings ----------------------------------------------------------------
    _add_heading(doc, "7. Audit Findings")
    audit_rows = [
        ("Critical findings", _metric_by_id(result, "critical-findings")),
        ("Open audit findings", _metric_by_id(result, "open-findings")),
        ("Overdue audit findings", _metric_by_id(result, "overdue-findings")),
    ]
    _add_table(
        doc,
        ["Audit metric", "Value"],
        [
            [label, str(value) if value is not None else "n/a"]
            for label, value in audit_rows
        ],
    )
    audit_docs = _documents_of_category(result, "audit-findings", "gia-findings")
    if audit_docs:
        doc.add_paragraph(
            "Aggregate audit figures are derived from the uploaded audit finding document(s): "
            + ", ".join(_clean(doc_.filename) for doc_ in audit_docs)
            + "."
        )

    # --- Exceptions --------------------------------------------------------------------
    _add_heading(doc, "8. Exceptions")
    exception_rows = [
        ("Open exceptions", _metric_by_id(result, "open-exceptions")),
        ("Overdue exceptions", _metric_by_id(result, "overdue-exceptions")),
    ]
    _add_table(
        doc,
        ["Exception metric", "Value"],
        [
            [label, str(value) if value is not None else "n/a"]
            for label, value in exception_rows
        ],
    )
    exception_docs = _documents_of_category(result, "exception-log")
    if exception_docs:
        doc.add_paragraph(
            "Aggregate exception figures are derived from the uploaded exception log document(s): "
            + ", ".join(_clean(doc_.filename) for doc_ in exception_docs)
            + "."
        )

    # --- AI recommendations ------------------------------------------------------------
    _add_heading(doc, "9. AI Recommendations")
    recommendations = result.recommendations or []
    if not recommendations:
        doc.add_paragraph("No AI recommendations were generated for this analysis run.")
    else:
        _add_table(
            doc,
            ["Priority", "Category", "Recommended Action", "Impact", "Confidence", "Evidence"],
            [
                [
                    recommendation.priority,
                    recommendation.category,
                    recommendation.action,
                    recommendation.impact,
                    f"{recommendation.confidence}%",
                    _evidence_text(recommendation.evidence),
                ]
                for recommendation in recommendations
            ],
        )

    # --- Management actions ------------------------------------------------------------
    _add_heading(doc, "10. Management Actions")
    actions = result.managementActions or []
    if not actions:
        doc.add_paragraph("No management actions were recorded for this analysis run.")
    else:
        _add_table(
            doc,
            ["ID", "Action", "Owner", "Department", "Due Date", "Priority", "Status"],
            [
                [
                    action.id,
                    action.action,
                    action.owner,
                    action.department,
                    action.dueDate,
                    action.priority,
                    action.status,
                ]
                for action in actions
            ],
        )

    # --- Sources & evidence ------------------------------------------------------------
    _add_heading(doc, "11. Sources and Evidence")
    documents = result.documents or []
    if not documents:
        doc.add_paragraph("No source documents were included in this analysis run.")
    else:
        _add_table(
            doc,
            ["Filename", "Category", "Status", "Evidence References"],
            [
                [
                    doc_.filename,
                    doc_.category,
                    doc_.status,
                    str(doc_.evidenceCount),
                ]
                for doc_ in documents
            ],
        )
    doc.add_paragraph(
        "Evidence references cite the exact row or section of the source "
        "document from which each finding or recommendation was derived."
    )

    # --- Confidence, warnings, disclaimer ----------------------------------------------
    _add_heading(doc, "12. Confidence and Disclaimer")
    doc.add_paragraph(
        f"Overall confidence in this analysis: {result.confidence}%. "
        "Confidence reflects data completeness, source coverage and the "
        "agreement of the underlying indicators."
    )
    if result.warnings:
        doc.add_paragraph("Warnings:").runs[0].bold = True
        for warning in result.warnings:
            doc.add_paragraph(warning, style="List Bullet")
    doc.add_paragraph(
        "This report was generated by Risk Copilot AI from the uploaded source "
        "documents and is provided for management information. It does not "
        "constitute legal or regulatory advice and should be reviewed by the "
        "operational risk team before distribution."
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
