"""Generate the Meridian Bank fictional operational-risk demo dataset.

Produces the five files the analysis pipeline ingests:
  Risk Register.xlsx        - 38 risks across 8 divisions
  Audit Findings.xlsx       - 22 audit findings (GIA / Internal Audit)
  Exception Log.xlsx        - 14 policy exceptions
  MIS.xlsx                  - 12 months of KRI history (2025-08 .. 2026-07)
  Operational Risk Policy.docx - short fictional policy (grounding material)

The data is a single coherent story:
  * the final MIS month (2026-07) matches the Risk Register population,
  * audit findings map to the same themes as the top risks,
  * exceptions reinforce those themes,
  * remediation drives a gentle downward trend across the year.

Column names are chosen to be recognised by the existing header matchers in
backend/app/analysis/pipeline/headers.py. No analysis logic is modified.

Run from the repository root:
    python sample-data/generate_demo_data.py
"""

from __future__ import annotations

from datetime import date, timedelta

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

# Fictional reference date for the July/August 2026 reporting cycle. Exception
# "days open" values are derived from this date so the columns stay consistent.
REPORT_DATE = date(2026, 8, 6)

OUT_DIR = "sample-data"

HEADER_FILL = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
HEADER_FONT = Font(color="FFFFFF", bold=True)
HEADER_ALIGN = Alignment(horizontal="center", vertical="center")


def _write_sheet(workbook: Workbook, title: str, headers: list[str], rows: list[list]) -> None:
    sheet = workbook.active
    sheet.title = title
    sheet.append(headers)
    for column_index in range(1, len(headers) + 1):
        cell = sheet.cell(row=1, column=column_index)
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = HEADER_ALIGN
    for row in rows:
        sheet.append(row)
    for column_index, _ in enumerate(headers, start=1):
        sheet.column_dimensions[get_column_letter(column_index)].width = 42
    sheet.freeze_panes = "A2"


# ---------------------------------------------------------------------------
# 1. Risk Register - 38 risks
# ---------------------------------------------------------------------------

RISK_HEADERS = [
    "Risk ID",
    "Risk Description",
    "Risk Category",
    "Division",
    "Likelihood",
    "Impact",
    "Inherent Risk",
    "Residual Risk",
    "Risk Owner",
    "Mitigation",
    "Status",
    "Review Date",
]

# (id, description, category, division, likelihood, impact, inherent, residual,
#  owner, mitigation, status, review_date)
RISKS = [
    ("RISK-2026-001",
     "Privileged access to production systems is not revoked on role change, creating an insider threat on core banking platforms",
     "Cybersecurity / Tech", "Technology", 5, 5, "Critical", "Critical",
     "CISO / IT Security",
     "Just-In-Time privileged access management rollout; quarterly access recertification",
     "Open", date(2026, 8, 31)),
    ("RISK-2026-002",
     "Concentration on a single cloud provider exposes core online banking to prolonged outage",
     "Third-Party / Vendor", "Operations", 4, 5, "Critical", "Critical",
     "Head of Cloud Operations",
     "Dual-cloud resilience pilot; contractual SLA credits; annual exit-plan review",
     "Open", date(2026, 9, 30)),
    ("RISK-2026-003",
     "Payment processing controls can fail during high-volume windows, allowing duplicate or unauthorized payments",
     "Payment Operations", "Payments", 4, 5, "Critical", "Critical",
     "Head of Payments",
     "Automated positive-pay reconciliation; four-eye review above thresholds; real-time anomaly alerts",
     "In Progress", date(2026, 8, 31)),
    ("RISK-2026-004",
     "KYC documentation for high-risk clients is incomplete, increasing money-laundering exposure",
     "AML / Compliance", "Compliance", 5, 4, "Critical", "Critical",
     "Chief Compliance Officer",
     "Enhanced due diligence workflow; automated document-expiry tracking; monthly file quality review",
     "In Progress", date(2026, 9, 30)),
    ("RISK-2026-005",
     "Segregation of duties gaps allow a single officer to both book and approve retail credit transactions",
     "Operations Risk", "Retail Banking", 4, 4, "High", "High",
     "Head of Retail Credit",
     "System-enforced dual authorisation; quarterly segregation-of-duties rule review",
     "Open", date(2026, 8, 31)),
    ("RISK-2026-006",
     "Unauthorised access to sensitive customer data by internal users holding broad file permissions",
     "Data Management", "Technology", 3, 5, "High", "High",
     "CISO / IT Security",
     "Attribute-based access control; data loss prevention monitoring; least-privilege cleanup",
     "In Progress", date(2026, 10, 31)),
    ("RISK-2026-007",
     "Reconciliation exceptions in Treasury remain unresolved past SLA, masking funding mismatches",
     "Operations Risk", "Treasury", 4, 4, "High", "High",
     "Treasury Operations Lead",
     "Automated reconciliation engine; SLA escalation to Deputy Treasurer",
     "Overdue", date(2026, 7, 31)),
    ("RISK-2026-008",
     "Third-party vendors hold broad data access rights beyond contractual need",
     "Third-Party / Vendor", "Operations", 3, 5, "High", "High",
     "Vendor Risk Manager",
     "Contractual access schedules; quarterly vendor access attestation",
     "In Progress", date(2026, 9, 30)),
    ("RISK-2026-009",
     "Loan officers can override credit limits without secondary approval, enabling internal fraud",
     "Fraud Risk", "Corporate Banking", 3, 5, "High", "High",
     "Head of Corporate Credit",
     "Override alerting to Chief Risk Officer; maker-checker on all limit exceptions",
     "In Progress", date(2026, 8, 31)),
    ("RISK-2026-010",
     "Ransomware attack on internet banking could disrupt retail services and leak credentials",
     "Cybersecurity / Tech", "Technology", 4, 4, "High", "High",
     "CISO / IT Security",
     "Immutable backups; annual red-team testing; network micro-segmentation",
     "In Progress", date(2026, 12, 31)),
    ("RISK-2026-011",
     "Credit model validation backlog may embed stale assumptions in capital allocations",
     "Model Risk", "Corporate Banking", 4, 4, "High", "High",
     "Chief Credit Officer",
     "Independent model validation plan; annual recalibration",
     "Open", date(2026, 11, 30)),
    ("RISK-2026-012",
     "Transaction monitoring rules are not updated for emerging scam typologies",
     "AML / Compliance", "Compliance", 4, 4, "High", "High",
     "Head of Financial Crime",
     "Quarterly typology review; machine-learning monitoring pilots",
     "Open", date(2026, 9, 30)),
    ("RISK-2026-013",
     "Social-engineering payment fraud targets staff through invoice redirect scams",
     "Fraud Risk", "Payments", 4, 4, "High", "High",
     "Head of Payments",
     "Payment fraud e-learning; callback verification for changes to payee details",
     "In Progress", date(2026, 8, 31)),
    ("RISK-2026-014",
     "All critical applications run from a single data centre, threatening business continuity",
     "Business Continuity", "Operations", 3, 5, "High", "High",
     "Head of Business Continuity",
     "Secondary site failover testing; RTO / RPO dashboard",
     "Open", date(2026, 12, 31)),
    ("RISK-2026-015",
     "Sanctions screening lists are applied inconsistently across product onboarding",
     "Legal / Regulatory", "Compliance", 4, 4, "High", "High",
     "Head of Sanctions",
     "Centralised screening engine; name-matching tuning; QA sampling",
     "In Progress", date(2026, 10, 31)),
    ("RISK-2026-016",
     "Key wealth-management processes depend on a small number of specialists, creating single-person dependency",
     "People / Human Capital", "Wealth Management", 4, 3, "High", "High",
     "Head of Wealth Management",
     "Succession planning; cross-training in advisory operations",
     "Open", date(2026, 9, 30)),
    ("RISK-2026-017",
     "Outsourced retail lending operations lack clear performance and incident oversight",
     "Third-Party / Vendor", "Retail Banking", 4, 4, "High", "High",
     "Head of Retail Operations",
     "Vendor scorecards; monthly service-level reviews",
     "In Progress", date(2026, 10, 31)),
    ("RISK-2026-018",
     "Payment settlement liquidity risk emerges when intraday limits are breached late in the day",
     "Payment Operations", "Treasury", 4, 4, "High", "High",
     "Treasury Operations Lead",
     "Intraday liquidity monitoring; pre-funded settlement buffers",
     "Open", date(2026, 8, 31)),
    ("RISK-2026-019",
     "IT change management gaps allow poorly tested changes into production",
     "Cybersecurity / Tech", "Technology", 4, 3, "High", "High",
     "Head of Technology Operations",
     "Change advisory board; automated regression testing; change freeze during peak periods",
     "In Progress", date(2026, 8, 31)),
    ("RISK-2026-020",
     "Manual data entry errors in retail lending applications produce inaccurate customer records",
     "Operations Risk", "Retail Banking", 4, 3, "Medium", "Medium",
     "Head of Retail Operations",
     "Input validation rules; two-person review for large applications",
     "Closed", date(2026, 7, 31)),
    ("RISK-2026-021",
     "Mobile banking application availability degrades during promotional traffic spikes",
     "Business Continuity", "Retail Banking", 3, 4, "Medium", "Medium",
     "Head of Digital Banking",
     "Auto-scaling; load testing before marketing campaigns",
     "In Progress", date(2026, 9, 30)),
    ("RISK-2026-022",
     "Endpoint patching lags behind vendor releases, widening the vulnerability window",
     "Cybersecurity / Tech", "Technology", 3, 4, "Medium", "Medium",
     "Head of Technology Operations",
     "Automated patch management; critical-patch SLA of 7 days",
     "In Progress", date(2026, 8, 31)),
    ("RISK-2026-023",
     "Trade finance documents are not consistently verified against shipping evidence",
     "AML / Compliance", "Corporate Banking", 3, 4, "Medium", "Medium",
     "Head of Trade Operations",
     "Digital document verification; checklist sign-off",
     "Open", date(2026, 10, 31)),
    ("RISK-2026-024",
     "Wealth suitability assessments are not refreshed after material client life events",
     "Legal / Regulatory", "Wealth Management", 3, 4, "Medium", "Medium",
     "Head of Wealth Management",
     "Client-review triggers; automated portfolio-health checks",
     "In Progress", date(2026, 9, 30)),
    ("RISK-2026-025",
     "Card fraud detection rules lag new merchant categories, delaying blocking",
     "Fraud Risk", "Payments", 3, 4, "Medium", "Medium",
     "Head of Payments",
     "Real-time rules engine; fraud typology sprint reviews",
     "In Progress", date(2026, 8, 31)),
    ("RISK-2026-026",
     "Collateral valuations for corporate facilities are refreshed less frequently than policy requires",
     "Credit Operations", "Corporate Banking", 3, 3, "Medium", "Medium",
     "Collateral Manager",
     "Valuation calendar; automated revaluation reminders",
     "Open", date(2026, 11, 30)),
    ("RISK-2026-027",
     "Interest-rate risk reporting relies on spreadsheets prone to manual error",
     "Model Risk", "Treasury", 3, 3, "Medium", "Medium",
     "Treasury Controller",
     "System-generated IRRBB reports; monthly variance checks",
     "Closed", date(2026, 7, 31)),
    ("RISK-2026-028",
     "Cash handling discrepancies occur across retail branches during peak settlement",
     "Operations Risk", "Operations", 4, 2, "Medium", "Medium",
     "Head of Branch Operations",
     "Cash-in-transit reconciliations; CCTV review of exceptions",
     "In Progress", date(2026, 8, 31)),
    ("RISK-2026-029",
     "Vendor service-desk staff retain production access after project completion",
     "Third-Party / Vendor", "Technology", 3, 3, "Medium", "Medium",
     "Vendor Risk Manager",
     "Project-access expiry controls; quarterly access reviews",
     "Open", date(2026, 9, 30)),
    ("RISK-2026-030",
     "Regulatory reporting is sometimes submitted close to deadline, risking late filings",
     "Legal / Regulatory", "Compliance", 3, 3, "Medium", "Medium",
     "Head of Regulatory Reporting",
     "Automated filing calendar; peer review before submission",
     "In Progress", date(2026, 8, 31)),
    ("RISK-2026-031",
     "Branch network depends on local servers with limited resilience",
     "Business Continuity", "Retail Banking", 3, 3, "Medium", "Medium",
     "Head of Branch Operations",
     "Branch-to-cloud connectivity plan; UPS coverage",
     "Open", date(2026, 12, 31)),
    ("RISK-2026-032",
     "Loan documentation is not consistently captured in the central repository",
     "Data Management", "Operations", 4, 2, "Medium", "Medium",
     "Records Manager",
     "Imaging workflow; completeness checks at booking",
     "In Progress", date(2026, 9, 30)),
    ("RISK-2026-033",
     "Fraudulent account opening attempts bypass some front-line verification checks",
     "Fraud Risk", "Retail Banking", 3, 3, "Medium", "Medium",
     "Head of Retail Operations",
     "Biometric verification; fraud-prevention screening at onboarding",
     "In Progress", date(2026, 8, 31)),
    ("RISK-2026-034",
     "Physical records are disposed of without consistent shredding controls",
     "Data Management", "Operations", 2, 3, "Low", "Low",
     "Records Manager",
     "Locked shredding bins; monthly disposal audits",
     "Closed", date(2026, 6, 30)),
    ("RISK-2026-035",
     "Legacy Excel macros in Treasury lack version control and testing",
     "Cybersecurity / Tech", "Treasury", 2, 2, "Low", "Low",
     "Treasury Controller",
     "Macro inventory; migration to governed reporting platform",
     "In Progress", date(2026, 9, 30)),
    ("RISK-2026-036",
     "Compliance training completion lags for new joiners in front-office roles",
     "People / Human Capital", "Compliance", 3, 2, "Low", "Low",
     "Head of Learning & Development",
     "Automated onboarding training; completion dashboards",
     "In Progress", date(2026, 8, 31)),
    ("RISK-2026-037",
     "Wealth procedure manuals are not updated after system changes",
     "Operations Risk", "Wealth Management", 2, 2, "Low", "Low",
     "Head of Wealth Management",
     "Document-control workflow; annual procedure reviews",
     "Open", date(2026, 10, 31)),
    ("RISK-2026-038",
     "Some access recertification campaigns miss dormant accounts",
     "Cybersecurity / Tech", "Retail Banking", 2, 2, "Low", "Low",
     "CISO / IT Security",
     "Automated dormant-account detection; closed-loop remediation",
     "In Progress", date(2026, 8, 31)),
]

RISK_ROWS = [
    list(row) for row in RISKS
]


# ---------------------------------------------------------------------------
# 2. Audit Findings - 22 findings
# ---------------------------------------------------------------------------

AUDIT_HEADERS = [
    "Finding ID",
    "Finding",
    "Area",
    "Division",
    "Rating",
    "Status",
    "Due Date",
    "Owner",
    "Audit Source",
]

# (id, finding, area, division, rating, status, due_date, owner, audit_source)
# "Area" and "Division" intentionally carry the same business division so the
# header matcher uses them interchangeably without ambiguity.
AUDIT_FINDINGS = [
    ("AUD-2026-001",
     "Excessive privileged access to production databases not recertified for 18 months",
     "Technology", "Technology", "Critical", "Open", date(2026, 8, 31),
     "IT Infrastructure Head", "GIA"),
    ("AUD-2026-002",
     "KYC files for high-risk clients missing enhanced due diligence evidence",
     "Compliance", "Compliance", "Critical", "In Progress", date(2026, 9, 30),
     "Compliance Lead", "GIA"),
    ("AUD-2026-003",
     "Duplicate payment controls ineffective during high-volume windows",
     "Payments", "Payments", "Critical", "Overdue", date(2026, 7, 15),
     "Head of Payments", "GIA"),
    ("AUD-2026-004",
     "Core online banking depends on a single cloud provider with no validated exit plan",
     "Operations", "Operations", "High", "In Progress", date(2026, 9, 30),
     "Head of Cloud Operations", "GIA"),
    ("AUD-2026-005",
     "Segregation of duties conflicts identified in retail credit booking and approval",
     "Retail Banking", "Retail Banking", "High", "Open", date(2026, 8, 31),
     "Head of Retail Credit", "Internal Audit"),
    ("AUD-2026-006",
     "Treasury reconciliation exceptions unresolved beyond 30 days",
     "Treasury", "Treasury", "High", "Overdue", date(2026, 6, 30),
     "Treasury Operations Lead", "GIA"),
    ("AUD-2026-007",
     "Third-party vendors hold broader data access than contractual schedules permit",
     "Operations", "Operations", "High", "In Progress", date(2026, 9, 30),
     "CISO / IT Security", "Internal Audit"),
    ("AUD-2026-008",
     "Transaction monitoring rules not refreshed for current scam typologies",
     "Compliance", "Compliance", "High", "In Progress", date(2026, 10, 31),
     "Head of Financial Crime", "GIA"),
    ("AUD-2026-009",
     "Credit model validation backlog exceeds acceptable review cycles",
     "Corporate Banking", "Corporate Banking", "High", "Open", date(2026, 11, 30),
     "Chief Credit Officer", "Internal Audit"),
    ("AUD-2026-010",
     "Ransomware recovery testing incomplete for internet banking components",
     "Technology", "Technology", "High", "Overdue", date(2026, 7, 31),
     "IT Infrastructure Head", "GIA"),
    ("AUD-2026-011",
     "Manual data entry controls in retail lending produce unresolved data-quality issues",
     "Retail Banking", "Retail Banking", "Medium", "In Progress", date(2026, 9, 30),
     "Head of Retail Operations", "Internal Audit"),
    ("AUD-2026-012",
     "Mobile banking availability monitoring does not cover all critical user journeys",
     "Retail Banking", "Retail Banking", "Medium", "Open", date(2026, 10, 31),
     "Head of Digital Banking", "Internal Audit"),
    ("AUD-2026-013",
     "Endpoint patching backlog exceeds internal targets on a quarter of devices",
     "Technology", "Technology", "Medium", "In Progress", date(2026, 8, 31),
     "IT Infrastructure Head", "Internal Audit"),
    ("AUD-2026-014",
     "Trade finance document verification gaps persist in high-value transactions",
     "Corporate Banking", "Corporate Banking", "Medium", "Open", date(2026, 11, 30),
     "Head of Trade Operations", "Internal Audit"),
    ("AUD-2026-015",
     "Wealth suitability assessments not refreshed after material client life events",
     "Wealth Management", "Wealth Management", "Medium", "In Progress", date(2026, 10, 31),
     "Head of Wealth Management", "Internal Audit"),
    ("AUD-2026-016",
     "Card fraud detection latency delays blocking of high-risk merchant categories",
     "Payments", "Payments", "Medium", "In Progress", date(2026, 8, 31),
     "Head of Payments", "GIA"),
    ("AUD-2026-017",
     "Collateral revaluation frequency below policy for commercial facilities",
     "Corporate Banking", "Corporate Banking", "Medium", "Open", date(2026, 12, 31),
     "Collateral Manager", "Internal Audit"),
    ("AUD-2026-018",
     "Regulatory reporting timelines at risk of non-compliance on peak filing days",
     "Compliance", "Compliance", "Medium", "In Progress", date(2026, 9, 30),
     "Head of Regulatory Reporting", "Internal Audit"),
    ("AUD-2026-019",
     "Loan record retention schedules not consistently enforced across branches",
     "Operations", "Operations", "Low", "Closed", date(2026, 5, 31),
     "Records Manager", "Internal Audit"),
    ("AUD-2026-020",
     "Branch resilience testing gaps identified in disaster-recovery exercise",
     "Retail Banking", "Retail Banking", "Low", "Closed", date(2026, 6, 30),
     "Head of Branch Operations", "Internal Audit"),
    ("AUD-2026-021",
     "Excel-based reporting controls in Treasury lack version governance",
     "Treasury", "Treasury", "Low", "Closed", date(2026, 4, 30),
     "Treasury Controller", "Internal Audit"),
    ("AUD-2026-022",
     "Access recertification campaigns occasionally miss dormant accounts",
     "Retail Banking", "Retail Banking", "Low", "Closed", date(2026, 5, 31),
     "Access Management Lead", "Internal Audit"),
]

AUDIT_ROWS = [
    list(row) for row in AUDIT_FINDINGS
]


# ---------------------------------------------------------------------------
# 3. Exception Log - 14 exceptions
# ---------------------------------------------------------------------------

EXCEPTION_HEADERS = [
    "Exception ID",
    "Exception Description",
    "Division",
    "Severity",
    "Status",
    "Raised Date",
    "Days Open",
    "Owner",
]

# (id, description, division, severity, status, days_open, owner)
# Days open is derived against REPORT_DATE to keep "Raised Date" consistent.
EXCEPTIONS = [
    ("EXC-2026-101",
     "Temporary extension of vendor access beyond policy while cloud migration completes",
     "Operations", "High", "Open", 78, "Vendor Risk Manager"),
    ("EXC-2026-102",
     "High-risk client onboarded pending completion of enhanced due diligence documentation",
     "Compliance", "Critical", "Open", 144, "Compliance Lead"),
    ("EXC-2026-103",
     "Temporary limit override on a corporate lending facility awaiting credit committee approval",
     "Corporate Banking", "Medium", "Open", 57, "Head of Corporate Credit"),
    ("EXC-2026-104",
     "Duplicate payment recovery not reversed within the standard three-day SLA",
     "Payments", "High", "Open", 35, "Head of Payments"),
    ("EXC-2026-105",
     "Segregation-of-duties exception allowing loan booking and approval by one officer during leave cover",
     "Retail Banking", "High", "In Progress", 19, "Head of Retail Credit"),
    ("EXC-2026-106",
     "Treasury reconciliation aging above threshold pending system reconciliation fix",
     "Treasury", "High", "Overdue", 106, "Treasury Operations Lead"),
    ("EXC-2026-107",
     "Privileged account password expiry deferral granted for a legacy production service",
     "Technology", "Medium", "Open", 12, "IT Infrastructure Head"),
    ("EXC-2026-108",
     "Out-of-hours third-party maintenance access approved for a critical network upgrade",
     "Technology", "Medium", "In Progress", 8, "Head of Technology Operations"),
    ("EXC-2026-109",
     "Card scheme rule deviation for recurring billing to a legacy merchant group",
     "Payments", "Medium", "Open", 17, "Head of Payments"),
    ("EXC-2026-110",
     "Cash variance sign-off delayed pending branch manager approval",
     "Operations", "Low", "In Progress", 9, "Head of Branch Operations"),
    ("EXC-2026-111",
     "Remote access for an offshore vendor support team extended beyond standard tenure",
     "Operations", "Medium", "Open", 41, "CISO / IT Security"),
    ("EXC-2026-112",
     "Wealth suitability assessment deferred for a client with recent adviser transition",
     "Wealth Management", "Medium", "Open", 22, "Head of Wealth Management"),
    ("EXC-2026-113",
     "Regulatory submission deadline waiver granted for a low-priority periodic return",
     "Compliance", "Low", "Closed", 68, "Head of Regulatory Reporting"),
    ("EXC-2026-114",
     "Minor retail branch process deviation for cash deposits above counter limit",
     "Retail Banking", "Low", "Closed", 62, "Head of Branch Operations"),
]

EXCEPTION_ROWS = [
    [
        item_id,
        description,
        division,
        severity,
        status,
        REPORT_DATE - timedelta(days=days_open),
        days_open,
        owner,
    ]
    for item_id, description, division, severity, status, days_open, owner in EXCEPTIONS
]


# ---------------------------------------------------------------------------
# 4. MIS - 12 months of KRI history
# ---------------------------------------------------------------------------

MIS_HEADERS = ["Metric ID", "Indicator", "Month", "Value", "Unit"]

# (month, critical, high, medium, low) risk counts.
# The final month (2026-07) matches the Risk Register population:
# Critical 4, High 15, Medium 14, Low 5.
MIS_HISTORY = [
    ("2025-08", 6, 21, 20, 10),
    ("2025-09", 6, 20, 19, 10),
    ("2025-10", 5, 20, 19, 9),
    ("2025-11", 5, 19, 18, 9),
    ("2025-12", 5, 18, 18, 8),
    ("2026-01", 5, 18, 17, 8),
    ("2026-02", 5, 17, 17, 8),
    ("2026-03", 4, 17, 16, 7),
    ("2026-04", 4, 16, 16, 7),
    ("2026-05", 4, 16, 15, 6),
    ("2026-06", 4, 16, 15, 6),
    ("2026-07", 4, 15, 14, 5),
]

_INDICATORS = (
    ("critical", "Critical Risk Count"),
    ("high", "High Risk Count"),
    ("medium", "Medium Risk Count"),
    ("low", "Low Risk Count"),
)

MIS_ROWS: list[list] = []
for index, (month, critical, high, medium, low) in enumerate(MIS_HISTORY):
    values = {"critical": critical, "high": high, "medium": medium, "low": low}
    for key, indicator in _INDICATORS:
        MIS_ROWS.append(
            [f"KRI-RC-{index + 1:02d}-{key[:3].upper()}", indicator, month, values[key], "count"]
        )


def generate_risk_register() -> None:
    workbook = Workbook()
    _write_sheet(workbook, "Risk Register", RISK_HEADERS, RISK_ROWS)
    workbook.save(f"{OUT_DIR}/Risk Register.xlsx")


def generate_audit_findings() -> None:
    workbook = Workbook()
    _write_sheet(workbook, "Audit Findings", AUDIT_HEADERS, AUDIT_ROWS)
    workbook.save(f"{OUT_DIR}/Audit Findings.xlsx")


def generate_exception_log() -> None:
    workbook = Workbook()
    _write_sheet(workbook, "Exception Log", EXCEPTION_HEADERS, EXCEPTION_ROWS)
    workbook.save(f"{OUT_DIR}/Exception Log.xlsx")


def generate_mis() -> None:
    workbook = Workbook()
    _write_sheet(workbook, "Key Risk Indicators", MIS_HEADERS, MIS_ROWS)
    workbook.save(f"{OUT_DIR}/MIS.xlsx")


def generate_policy() -> None:
    document = Document()
    document.core_properties.title = "Operational Risk Management Policy"
    document.core_properties.author = "Meridian Bank Group Risk"

    document.add_heading("Operational Risk Management Policy", level=1)
    document.add_paragraph(
        "Version 4.2 | Issued by Group Risk Management | Meridian Bank (fictional entity)"
    )

    document.add_heading("1. Purpose", level=2)
    document.add_paragraph(
        "This policy defines the principles by which operational risk is identified, "
        "assessed, mitigated, monitored and reported across Meridian Bank. It applies to "
        "all divisions, subsidiaries and third-party service providers acting on the "
        "bank's behalf."
    )

    document.add_heading("2. Risk Appetite", level=2)
    document.add_paragraph(
        "Meridian Bank maintains a moderate-to-conservative operational risk appetite. "
        "The bank seeks to keep overall exposure within board-approved thresholds and to "
        "reduce the frequency and impact of high-severity operational losses over time."
    )

    document.add_heading("3. Governance", level=2)
    document.add_paragraph(
        "The Board Risk Committee sets risk appetite. The Group Risk function owns the "
        "operational risk framework, and each division designates a risk owner "
        "accountable for maintaining its risk register and control environment."
    )

    document.add_heading("4. Risk Identification and Assessment", level=2)
    document.add_paragraph(
        "Divisions maintain a current risk register covering inherent and residual "
        "exposure, likelihood and impact ratings, controls and remediation actions. "
        "Risks are reassessed at least annually or when material change occurs."
    )

    document.add_heading("5. Control and Mitigation", level=2)
    document.add_paragraph(
        "Controls must be documented and tested for effectiveness. Where residual risk "
        "remains above appetite, compensating controls or an approved exception are "
        "required. Exceptions are time-limited and tracked to closure."
    )

    document.add_heading("6. Monitoring and Reporting", level=2)
    document.add_paragraph(
        "Key risk indicators are monitored monthly against thresholds and reported to "
        "the Risk Committee. Significant incidents and audit findings are escalated "
        "within agreed timelines."
    )

    document.add_heading("7. Review Cycle", level=2)
    document.add_paragraph(
        "This policy is reviewed annually by Group Risk Management and approved by the "
        "Board Risk Committee."
    )

    document.save(f"{OUT_DIR}/Operational Risk Policy.docx")


def main() -> None:
    generate_risk_register()
    generate_audit_findings()
    generate_exception_log()
    generate_mis()
    generate_policy()
    print("Generated:")
    for filename in [
        "Risk Register.xlsx",
        "Audit Findings.xlsx",
        "Exception Log.xlsx",
        "MIS.xlsx",
        "Operational Risk Policy.docx",
    ]:
        print(f"  {filename}")
    from collections import Counter

    distribution = Counter(row[7] for row in RISK_ROWS)
    print(f"Risk distribution: {dict(distribution)}")
    print(f"Risk rows: {len(RISK_ROWS)}")
    print(f"Audit rows: {len(AUDIT_ROWS)}")
    print(f"Exception rows: {len(EXCEPTION_ROWS)}")
    print(f"MIS rows: {len(MIS_ROWS)}")


if __name__ == "__main__":
    main()
